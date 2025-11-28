#!/usr/bin/env python
"""
Script to create professional certificate templates for different ISO standards.
These templates use placeholders that will be replaced with actual data during generation.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def add_page_border(section):
    """Add a border to the page."""
    sectPr = section._sectPr
    pgBorders = OxmlElement('w:pgBorders')
    pgBorders.set(qn('w:offsetFrom'), 'page')
    
    for border_name in ('top', 'left', 'bottom', 'right'):
        border_el = OxmlElement(f'w:{border_name}')
        border_el.set(qn('w:val'), 'single')
        border_el.set(qn('w:sz'), '24')
        border_el.set(qn('w:space'), '24')
        border_el.set(qn('w:color'), '1F4788')
        pgBorders.append(border_el)
    
    sectPr.append(pgBorders)

def create_iso_9001_template():
    """Create ISO 9001:2015 Quality Management certificate template."""
    doc = Document()
    
    # Set margins
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.5)
    section.right_margin = Inches(1.5)
    
    # Add page border
    add_page_border(section)
    
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('CERTIFICATE OF REGISTRATION')
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(31, 71, 136)
    
    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Quality Management System')
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(31, 71, 136)
    
    doc.add_paragraph()  # Spacing
    
    # Certificate body
    body = doc.add_paragraph()
    body.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = body.add_run('This is to certify that')
    run.font.size = Pt(12)
    
    # Client name
    client_para = doc.add_paragraph()
    client_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = client_para.add_run('{{client_name}}')
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    
    # Address
    address_para = doc.add_paragraph()
    address_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = address_para.add_run('{{client_address}}')
    run.font.size = Pt(11)
    run.font.italic = True
    
    doc.add_paragraph()  # Spacing
    
    # Certification text
    cert_text = doc.add_paragraph()
    cert_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cert_text.add_run('has implemented and maintains a Quality Management System which complies with the requirements of')
    run.font.size = Pt(12)
    
    # ISO Standard
    iso_para = doc.add_paragraph()
    iso_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = iso_para.add_run('{{iso_standard_code}} - {{iso_standard_name}}')
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(31, 71, 136)
    
    doc.add_paragraph()  # Spacing
    
    # Scope
    scope_label = doc.add_paragraph()
    scope_label.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = scope_label.add_run('Scope of Certification:')
    run.font.size = Pt(11)
    run.font.bold = True
    
    scope_para = doc.add_paragraph()
    scope_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = scope_para.add_run('{{scope}}')
    run.font.size = Pt(11)
    
    doc.add_paragraph()  # Spacing
    
    # Certificate details table
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Light Grid Accent 1'
    
    # Certificate Number
    table.rows[0].cells[0].text = 'Certificate Number:'
    table.rows[0].cells[1].text = '{{certificate_number}}'
    
    # Issue Date
    table.rows[1].cells[0].text = 'Issue Date:'
    table.rows[1].cells[1].text = '{{issue_date}}'
    
    # Expiry Date
    table.rows[2].cells[0].text = 'Expiry Date:'
    table.rows[2].cells[1].text = '{{expiry_date}}'
    
    # Certification Body
    table.rows[3].cells[0].text = 'Certification Body:'
    table.rows[3].cells[1].text = '{{certification_body}}'
    
    doc.add_paragraph()  # Spacing
    
    # Signature section
    sig_para = doc.add_paragraph()
    sig_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = sig_para.add_run('_________________________\n')
    run.font.size = Pt(10)
    run = sig_para.add_run('{{lead_auditor_name}}\n')
    run.font.size = Pt(10)
    run.font.bold = True
    run = sig_para.add_run('Lead Auditor\n')
    run.font.size = Pt(9)
    run = sig_para.add_run('Accreditation No: {{accreditation_number}}')
    run.font.size = Pt(9)
    
    # Footer
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_para.add_run('This certificate remains the property of {{certification_body}} and is subject to periodic surveillance audits.')
    run.font.size = Pt(8)
    run.font.italic = True
    run.font.color.rgb = RGBColor(128, 128, 128)
    
    return doc

def create_iso_14001_template():
    """Create ISO 14001:2015 Environmental Management certificate template."""
    doc = Document()

    # Set margins
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.5)
    section.right_margin = Inches(1.5)

    # Add page border
    add_page_border(section)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('CERTIFICATE OF REGISTRATION')
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(34, 139, 34)  # Forest Green

    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Environmental Management System')
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(34, 139, 34)

    doc.add_paragraph()  # Spacing

    # Certification text
    cert_text = doc.add_paragraph()
    cert_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cert_text.add_run('This is to certify that')
    run.font.size = Pt(12)

    # Client name
    client_para = doc.add_paragraph()
    client_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = client_para.add_run('{{client_name}}')
    run.font.size = Pt(16)
    run.font.bold = True

    # Client address
    address_para = doc.add_paragraph()
    address_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = address_para.add_run('{{client_address}}')
    run.font.size = Pt(11)

    doc.add_paragraph()  # Spacing

    # Certification text
    cert_text = doc.add_paragraph()
    cert_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cert_text.add_run('has implemented and maintains an Environmental Management System which complies with the requirements of')
    run.font.size = Pt(12)

    # ISO Standard
    iso_para = doc.add_paragraph()
    iso_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = iso_para.add_run('{{iso_standard_code}} - {{iso_standard_name}}')
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(34, 139, 34)

    doc.add_paragraph()  # Spacing

    # Scope
    scope_label = doc.add_paragraph()
    scope_label.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = scope_label.add_run('Scope of Certification:')
    run.font.size = Pt(11)
    run.font.bold = True

    scope_para = doc.add_paragraph()
    scope_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = scope_para.add_run('{{scope}}')
    run.font.size = Pt(11)

    doc.add_paragraph()  # Spacing

    # Certificate details table
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Light Grid Accent 1'

    details = [
        ('Certificate Number:', '{{certificate_number}}'),
        ('Issue Date:', '{{issue_date}}'),
        ('Expiry Date:', '{{expiry_date}}'),
        ('Accreditation Number:', '{{accreditation_number}}'),
    ]

    for i, (label, value) in enumerate(details):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = value

    doc.add_paragraph()  # Spacing

    # Signature section
    sig_para = doc.add_paragraph()
    sig_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = sig_para.add_run('Authorized Signatory:')
    run.font.size = Pt(11)
    run.font.bold = True

    sig_name = doc.add_paragraph()
    run = sig_name.add_run('{{lead_auditor_name}}')
    run.font.size = Pt(11)

    sig_title = doc.add_paragraph()
    run = sig_title.add_run('Lead Auditor')
    run.font.size = Pt(10)
    run.font.italic = True

    # Footer
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_para.add_run('This certificate remains the property of {{certification_body}} and is subject to periodic surveillance audits.')
    run.font.size = Pt(8)
    run.font.italic = True
    run.font.color.rgb = RGBColor(128, 128, 128)

    return doc


def create_iso_45001_template():
    """Create ISO 45001:2018 Occupational Health & Safety certificate template."""
    doc = Document()

    # Set margins
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.5)
    section.right_margin = Inches(1.5)

    # Add page border
    add_page_border(section)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('CERTIFICATE OF REGISTRATION')
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(220, 20, 60)  # Crimson

    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Occupational Health & Safety Management System')
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(220, 20, 60)

    doc.add_paragraph()  # Spacing

    # Certification text
    cert_text = doc.add_paragraph()
    cert_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cert_text.add_run('This is to certify that')
    run.font.size = Pt(12)

    # Client name
    client_para = doc.add_paragraph()
    client_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = client_para.add_run('{{client_name}}')
    run.font.size = Pt(16)
    run.font.bold = True

    # Client address
    address_para = doc.add_paragraph()
    address_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = address_para.add_run('{{client_address}}')
    run.font.size = Pt(11)

    doc.add_paragraph()  # Spacing

    # Certification text
    cert_text = doc.add_paragraph()
    cert_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cert_text.add_run('has implemented and maintains an Occupational Health & Safety Management System which complies with the requirements of')
    run.font.size = Pt(12)

    # ISO Standard
    iso_para = doc.add_paragraph()
    iso_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = iso_para.add_run('{{iso_standard_code}} - {{iso_standard_name}}')
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(220, 20, 60)

    doc.add_paragraph()  # Spacing

    # Scope
    scope_label = doc.add_paragraph()
    scope_label.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = scope_label.add_run('Scope of Certification:')
    run.font.size = Pt(11)
    run.font.bold = True

    scope_para = doc.add_paragraph()
    scope_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = scope_para.add_run('{{scope}}')
    run.font.size = Pt(11)

    doc.add_paragraph()  # Spacing

    # Certificate details table
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Light Grid Accent 1'

    details = [
        ('Certificate Number:', '{{certificate_number}}'),
        ('Issue Date:', '{{issue_date}}'),
        ('Expiry Date:', '{{expiry_date}}'),
        ('Accreditation Number:', '{{accreditation_number}}'),
    ]

    for i, (label, value) in enumerate(details):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = value

    doc.add_paragraph()  # Spacing

    # Signature section
    sig_para = doc.add_paragraph()
    sig_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = sig_para.add_run('Authorized Signatory:')
    run.font.size = Pt(11)
    run.font.bold = True

    sig_name = doc.add_paragraph()
    run = sig_name.add_run('{{lead_auditor_name}}')
    run.font.size = Pt(11)

    sig_title = doc.add_paragraph()
    run = sig_title.add_run('Lead Auditor')
    run.font.size = Pt(10)
    run.font.italic = True

    # Footer
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_para.add_run('This certificate remains the property of {{certification_body}} and is subject to periodic surveillance audits.')
    run.font.size = Pt(8)
    run.font.italic = True
    run.font.color.rgb = RGBColor(128, 128, 128)

    return doc


def create_iso_27001_template():
    """Create ISO 27001:2013 Information Security certificate template."""
    doc = Document()

    # Set margins
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.5)
    section.right_margin = Inches(1.5)

    # Add page border
    add_page_border(section)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('CERTIFICATE OF REGISTRATION')
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(70, 130, 180)  # Steel Blue

    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Information Security Management System')
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(70, 130, 180)

    doc.add_paragraph()  # Spacing

    # Certification text
    cert_text = doc.add_paragraph()
    cert_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cert_text.add_run('This is to certify that')
    run.font.size = Pt(12)

    # Client name
    client_para = doc.add_paragraph()
    client_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = client_para.add_run('{{client_name}}')
    run.font.size = Pt(16)
    run.font.bold = True

    # Client address
    address_para = doc.add_paragraph()
    address_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = address_para.add_run('{{client_address}}')
    run.font.size = Pt(11)

    doc.add_paragraph()  # Spacing

    # Certification text
    cert_text = doc.add_paragraph()
    cert_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cert_text.add_run('has implemented and maintains an Information Security Management System which complies with the requirements of')
    run.font.size = Pt(12)

    # ISO Standard
    iso_para = doc.add_paragraph()
    iso_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = iso_para.add_run('{{iso_standard_code}} - {{iso_standard_name}}')
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(70, 130, 180)

    doc.add_paragraph()  # Spacing

    # Scope
    scope_label = doc.add_paragraph()
    scope_label.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = scope_label.add_run('Scope of Certification:')
    run.font.size = Pt(11)
    run.font.bold = True

    scope_para = doc.add_paragraph()
    scope_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = scope_para.add_run('{{scope}}')
    run.font.size = Pt(11)

    doc.add_paragraph()  # Spacing

    # Certificate details table
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Light Grid Accent 1'

    details = [
        ('Certificate Number:', '{{certificate_number}}'),
        ('Issue Date:', '{{issue_date}}'),
        ('Expiry Date:', '{{expiry_date}}'),
        ('Accreditation Number:', '{{accreditation_number}}'),
    ]

    for i, (label, value) in enumerate(details):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = value

    doc.add_paragraph()  # Spacing

    # Signature section
    sig_para = doc.add_paragraph()
    sig_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = sig_para.add_run('Authorized Signatory:')
    run.font.size = Pt(11)
    run.font.bold = True

    sig_name = doc.add_paragraph()
    run = sig_name.add_run('{{lead_auditor_name}}')
    run.font.size = Pt(11)

    sig_title = doc.add_paragraph()
    run = sig_title.add_run('Lead Auditor')
    run.font.size = Pt(10)
    run.font.italic = True

    # Footer
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_para.add_run('This certificate remains the property of {{certification_body}} and is subject to periodic surveillance audits.')
    run.font.size = Pt(8)
    run.font.italic = True
    run.font.color.rgb = RGBColor(128, 128, 128)

    return doc


def create_iso_22301_template():
    """
    Create ISO 22301:2019 Business Continuity certificate template.

    Based on analysis of BUSINESS CONTINUITY_Stage_II_Certification_Audit_Report.docx
    Color scheme: Dark Goldenrod (RGB: 184, 134, 11)

    Field mappings from audit report:
    - Client Name → {{client_name}}
    - Address → {{client_address}}
    - Phone No. → {{client_phone}}
    - Email Address → {{client_email}}
    - Lead Auditor → {{lead_auditor_name}}
    - Audit Scope → {{scope}}
    - Conducted on → {{issue_date}}
    """
    doc = Document()

    # Set margins (A4 size)
    section = doc.sections[0]
    section.page_width = Inches(8.27)  # A4 width
    section.page_height = Inches(11.69)  # A4 height
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.5)
    section.right_margin = Inches(1.5)

    # Add page border
    add_page_border(section)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('CERTIFICATE OF REGISTRATION')
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(184, 134, 11)  # Dark Goldenrod

    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Business Continuity Management System')
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(184, 134, 11)

    doc.add_paragraph()  # Spacing

    # Certification text
    cert_text = doc.add_paragraph()
    cert_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cert_text.add_run('This is to certify that')
    run.font.size = Pt(12)

    # Client name
    client_para = doc.add_paragraph()
    client_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = client_para.add_run('{{client_name}}')
    run.font.size = Pt(16)
    run.font.bold = True

    # Client address
    address_para = doc.add_paragraph()
    address_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = address_para.add_run('{{client_address}}')
    run.font.size = Pt(11)

    # Client contact information (from audit report)
    contact_para = doc.add_paragraph()
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = contact_para.add_run('Phone: {{client_phone}} | Email: {{client_email}}')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()  # Spacing

    # Certification text
    cert_text = doc.add_paragraph()
    cert_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cert_text.add_run('has implemented and maintains a Business Continuity Management System\nwhich complies with the requirements of')
    run.font.size = Pt(12)

    # ISO Standard
    iso_para = doc.add_paragraph()
    iso_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = iso_para.add_run('{{iso_standard_code}} - {{iso_standard_name}}')
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(184, 134, 11)

    doc.add_paragraph()  # Spacing

    # Scope
    scope_label = doc.add_paragraph()
    scope_label.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = scope_label.add_run('Scope of Certification:')
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGBColor(184, 134, 11)

    scope_para = doc.add_paragraph()
    scope_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = scope_para.add_run('{{scope}}')
    run.font.size = Pt(11)

    doc.add_paragraph()  # Spacing

    # Certificate details table
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Light Grid Accent 1'

    details = [
        ('Certificate Number:', '{{certificate_number}}'),
        ('Issue Date:', '{{issue_date}}'),
        ('Expiry Date:', '{{expiry_date}}'),
        ('Certification Body:', '{{certification_body}}'),
        ('Accreditation Number:', '{{accreditation_number}}'),
    ]

    for i, (label, value) in enumerate(details):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = value
        # Bold the labels
        for paragraph in table.rows[i].cells[0].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    doc.add_paragraph()  # Spacing

    # Signature section
    sig_para = doc.add_paragraph()
    sig_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = sig_para.add_run('Authorized Signatory:')
    run.font.size = Pt(11)
    run.font.bold = True

    sig_name = doc.add_paragraph()
    run = sig_name.add_run('{{lead_auditor_name}}')
    run.font.size = Pt(11)

    sig_title = doc.add_paragraph()
    run = sig_title.add_run('Lead Auditor')
    run.font.size = Pt(10)
    run.font.italic = True

    # Footer
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_para.add_run('This certificate remains the property of {{certification_body}} and is subject to periodic surveillance audits.\nCertificate issued based on Stage II Certification Audit conducted in accordance with ISO 22301:2019 requirements.')
    run.font.size = Pt(8)
    run.font.italic = True
    run.font.color.rgb = RGBColor(128, 128, 128)

    return doc


def main():
    """Create all certificate templates."""
    output_dir = 'media/certificate_templates/samples'
    os.makedirs(output_dir, exist_ok=True)

    templates = [
        ('ISO 9001:2015', create_iso_9001_template, 'ISO_9001_2015_Certificate_Template.docx'),
        ('ISO 14001:2015', create_iso_14001_template, 'ISO_14001_2015_Certificate_Template.docx'),
        ('ISO 45001:2018', create_iso_45001_template, 'ISO_45001_2018_Certificate_Template.docx'),
        ('ISO 27001:2013', create_iso_27001_template, 'ISO_27001_2013_Certificate_Template.docx'),
        ('ISO 22301:2019', create_iso_22301_template, 'ISO_22301_2019_Certificate_Template.docx'),
    ]

    print("=" * 80)
    print("Creating ISO Certificate Templates")
    print("=" * 80)

    for name, create_func, filename in templates:
        print(f"\nCreating {name} certificate template...")
        doc = create_func()
        output_path = os.path.join(output_dir, filename)
        doc.save(output_path)
        print(f"✓ Created: {output_path}")

    print("\n" + "=" * 80)
    print("Template creation complete!")
    print("=" * 80)
    print("\nAvailable placeholders:")
    print("  - {{client_name}}")
    print("  - {{client_address}}")
    print("  - {{client_email}}")
    print("  - {{client_phone}}")
    print("  - {{iso_standard_code}}")
    print("  - {{iso_standard_name}}")
    print("  - {{certificate_number}}")
    print("  - {{issue_date}}")
    print("  - {{expiry_date}}")
    print("  - {{scope}}")
    print("  - {{certification_body}}")
    print("  - {{accreditation_number}}")
    print("  - {{lead_auditor_name}}")
    print("  - {{lead_auditor_email}}")

if __name__ == '__main__':
    main()

