"""
Certificate generation services.
Handles template processing and PDF generation.
"""
import os
from datetime import datetime
from django.conf import settings
from django.core.files.base import ContentFile
from .models import Certification, CertificateTemplate, CertificationHistory


class CertificateGenerationService:
    """
    Service for generating certificates from templates.
    Supports multiple template types: DOCX, HTML, PDF.
    """
    
    def __init__(self, certification: Certification):
        self.certification = certification
        self.template = certification.template
        
    def generate(self, user=None):
        """
        Generate certificate document from template.
        Returns the generated file path.
        """
        if not self.template:
            raise ValueError("No template assigned to this certification")
        
        if self.template.template_type == 'docx':
            return self._generate_docx()
        elif self.template.template_type == 'html':
            return self._generate_html_to_pdf()
        elif self.template.template_type == 'pdf':
            return self._generate_pdf()
        else:
            raise ValueError(f"Unsupported template type: {self.template.template_type}")
    
    def _get_template_variables(self):
        """
        Get all variables to populate in the template.
        Returns a dictionary of variable names and values.
        """
        cert = self.certification
        
        variables = {
            # Certificate Information
            'certificate_number': cert.certificate_number,
            'issue_date': cert.issue_date.strftime('%B %d, %Y'),
            'expiry_date': cert.expiry_date.strftime('%B %d, %Y'),
            'scope': cert.scope,
            
            # Client Information
            'client_name': cert.client.name,
            'client_address': cert.client.address,
            'client_email': cert.client.email,
            'client_phone': cert.client.phone,
            'client_industry': cert.client.industry or '',
            
            # ISO Standard
            'iso_standard_code': cert.iso_standard.code,
            'iso_standard_name': cert.iso_standard.name,
            'iso_standard_description': cert.iso_standard.description,
            
            # Auditor Information
            'lead_auditor_name': f"{cert.lead_auditor.first_name} {cert.lead_auditor.last_name}" if cert.lead_auditor else '',
            'lead_auditor_email': cert.lead_auditor.email if cert.lead_auditor else '',
            
            # Certification Body
            'certification_body': cert.certification_body or '',
            'accreditation_number': cert.accreditation_number or '',
            
            # Audit Information (if available)
            'audit_number': cert.audit.audit_number if cert.audit else '',
            'audit_type': cert.audit.get_audit_type_display() if cert.audit else '',
            
            # Current Date
            'current_date': datetime.now().strftime('%B %d, %Y'),
            'current_year': datetime.now().year,
        }
        
        # Add any custom metadata
        if cert.metadata:
            variables.update(cert.metadata)
        
        return variables
    
    def _generate_docx(self):
        """
        Generate certificate from DOCX template using python-docx.
        Note: Requires python-docx library to be installed.
        """
        try:
            from docx import Document
        except ImportError:
            raise ImportError("python-docx library is required for DOCX template generation. Install it with: pip install python-docx")
        
        # Load template
        template_path = self.template.template_file.path
        doc = Document(template_path)
        
        # Get variables
        variables = self._get_template_variables()
        
        # Replace placeholders in paragraphs
        for paragraph in doc.paragraphs:
            for key, value in variables.items():
                placeholder = f"{{{{{key}}}}}"  # {{variable_name}}
                if placeholder in paragraph.text:
                    paragraph.text = paragraph.text.replace(placeholder, str(value))
        
        # Replace placeholders in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, value in variables.items():
                        placeholder = f"{{{{{key}}}}}"
                        if placeholder in cell.text:
                            cell.text = cell.text.replace(placeholder, str(value))
        
        # Save generated document
        output_filename = f"{self.certification.certificate_number}.docx"
        output_path = os.path.join(
            settings.MEDIA_ROOT,
            'certificates',
            str(datetime.now().year),
            str(datetime.now().month).zfill(2),
            output_filename
        )
        
        # Create directory if it doesn't exist
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        
        # Save document
        doc.save(output_path)
        
        # Update certification with document URL
        relative_path = os.path.relpath(output_path, settings.MEDIA_ROOT)
        self.certification.document_url = relative_path
        self.certification.save()
        
        return output_path
    
    def _generate_html_to_pdf(self):
        """
        Generate PDF from HTML template.
        Note: Requires weasyprint or similar library.
        """
        # TODO: Implement HTML to PDF generation
        raise NotImplementedError("HTML to PDF generation not yet implemented")
    
    def _generate_pdf(self):
        """
        Generate PDF using JasperReports or similar.
        Note: Requires JasperReports integration.
        """
        # TODO: Implement JasperReports PDF generation
        raise NotImplementedError("JasperReports PDF generation not yet implemented")

