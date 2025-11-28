#!/usr/bin/env python
"""
Script to analyze ISO certificate templates and extract their structure.
"""

from docx import Document
import os
import json

def analyze_docx_structure(file_path):
    """Analyze a DOCX file and extract its structure."""
    try:
        doc = Document(file_path)
        
        structure = {
            'file_name': os.path.basename(file_path),
            'paragraphs_count': len(doc.paragraphs),
            'tables_count': len(doc.tables),
            'sections': [],
            'tables': [],
            'fields': set()
        }
        
        # Analyze paragraphs
        current_section = None
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue
                
            # Check if it's a heading (bold, larger font, or specific patterns)
            is_heading = False
            if para.runs:
                first_run = para.runs[0]
                if first_run.bold or (first_run.font.size and first_run.font.size.pt > 12):
                    is_heading = True
            
            # Look for section markers
            if any(marker in text.upper() for marker in ['SECTION', 'PART', 'CHAPTER', '1.', '2.', '3.', '4.', '5.']):
                if len(text) < 100:  # Likely a heading
                    is_heading = True
            
            if is_heading and len(text) < 150:
                current_section = text
                structure['sections'].append({
                    'index': i,
                    'title': text,
                    'content_preview': []
                })
            elif current_section and structure['sections']:
                # Add to current section's content
                if len(structure['sections'][-1]['content_preview']) < 3:
                    structure['sections'][-1]['content_preview'].append(text[:100])
            
            # Extract potential field names (looking for patterns like "Name:", "Date:", etc.)
            if ':' in text:
                parts = text.split(':')
                if len(parts[0]) < 50:  # Likely a field label
                    field_name = parts[0].strip()
                    structure['fields'].add(field_name)
        
        # Analyze tables
        for table_idx, table in enumerate(doc.tables):
            table_data = {
                'index': table_idx,
                'rows': len(table.rows),
                'cols': len(table.columns),
                'headers': [],
                'sample_data': []
            }
            
            # Get headers (first row)
            if table.rows:
                table_data['headers'] = [cell.text.strip() for cell in table.rows[0].cells]
                
                # Get sample data (first 2 data rows)
                for row_idx in range(1, min(3, len(table.rows))):
                    row_data = [cell.text.strip()[:50] for cell in table.rows[row_idx].cells]
                    table_data['sample_data'].append(row_data)
            
            structure['tables'].append(table_data)
        
        structure['fields'] = sorted(list(structure['fields']))
        
        return structure
        
    except Exception as e:
        return {'error': str(e), 'file_name': os.path.basename(file_path)}

def main():
    """Analyze all ISO templates."""
    iso_dir = '../iso'
    
    if not os.path.exists(iso_dir):
        print(f"ISO directory not found: {iso_dir}")
        return
    
    # Get all DOCX files
    docx_files = [f for f in os.listdir(iso_dir) if f.endswith('.docx')]
    
    print("=" * 80)
    print("ISO CERTIFICATE TEMPLATES ANALYSIS")
    print("=" * 80)
    print(f"\nFound {len(docx_files)} DOCX templates\n")
    
    all_structures = {}
    
    for docx_file in sorted(docx_files):
        file_path = os.path.join(iso_dir, docx_file)
        print(f"\n{'='*80}")
        print(f"Analyzing: {docx_file}")
        print('='*80)
        
        structure = analyze_docx_structure(file_path)
        all_structures[docx_file] = structure
        
        if 'error' in structure:
            print(f"ERROR: {structure['error']}")
            continue
        
        print(f"\nDocument Statistics:")
        print(f"  - Paragraphs: {structure['paragraphs_count']}")
        print(f"  - Tables: {structure['tables_count']}")
        print(f"  - Sections: {len(structure['sections'])}")
        print(f"  - Fields: {len(structure['fields'])}")
        
        if structure['sections']:
            print(f"\nSections Found:")
            for section in structure['sections'][:10]:  # Show first 10
                print(f"  [{section['index']}] {section['title']}")
        
        if structure['fields']:
            print(f"\nFields Found:")
            for field in structure['fields'][:20]:  # Show first 20
                print(f"  - {field}")
        
        if structure['tables']:
            print(f"\nTables Found:")
            for table in structure['tables'][:3]:  # Show first 3
                print(f"  Table {table['index']}: {table['rows']}x{table['cols']}")
                if table['headers']:
                    print(f"    Headers: {', '.join(table['headers'][:5])}")
    
    # Save detailed analysis to JSON
    output_file = 'iso_templates_analysis.json'
    with open(output_file, 'w') as f:
        json.dump(all_structures, f, indent=2, default=str)
    
    print(f"\n{'='*80}")
    print(f"Detailed analysis saved to: {output_file}")
    print('='*80)

if __name__ == '__main__':
    main()

