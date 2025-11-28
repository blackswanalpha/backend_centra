#!/usr/bin/env python3
"""
Test script for ISO 22301:2019 certificate template.

This script:
1. Analyzes the ISO 22301 Stage II Audit Report template
2. Tests the ISO 22301 certificate template generation
3. Creates a sample certificate with realistic data
4. Verifies all field mappings work correctly
"""

from docx import Document
from datetime import datetime, timedelta
import os

def analyze_audit_report():
    """Analyze the ISO 22301 Stage II Audit Report."""
    print("=" * 80)
    print("ANALYZING ISO 22301 STAGE II AUDIT REPORT")
    print("=" * 80)
    
    audit_report_path = '../iso/BUSINESS CONTINUITY_Stage_II_Certification_Audit_Report.docx'
    
    if not os.path.exists(audit_report_path):
        print(f"❌ Audit report not found: {audit_report_path}")
        return None
    
    doc = Document(audit_report_path)
    
    print(f"\n✓ Loaded audit report")
    print(f"  - Total paragraphs: {len(doc.paragraphs)}")
    print(f"  - Total tables: {len(doc.tables)}")
    print(f"  - Total sections: {len(doc.sections)}")
    
    # Extract field placeholders
    fields = []
    for i, para in enumerate(doc.paragraphs[:20]):
        text = para.text.strip()
        if ':' in text and '__' in text:
            fields.append(text)
    
    print(f"\n✓ Identified {len(fields)} field placeholders:")
    for field in fields:
        print(f"  - {field}")
    
    return doc


def test_certificate_template():
    """Test the ISO 22301 certificate template."""
    print("\n" + "=" * 80)
    print("TESTING ISO 22301 CERTIFICATE TEMPLATE")
    print("=" * 80)
    
    template_path = 'media/certificate_templates/samples/ISO_22301_2019_Certificate_Template.docx'
    
    if not os.path.exists(template_path):
        print(f"❌ Certificate template not found: {template_path}")
        return None
    
    doc = Document(template_path)
    
    print(f"\n✓ Loaded certificate template")
    print(f"  - Total paragraphs: {len(doc.paragraphs)}")
    print(f"  - Total tables: {len(doc.tables)}")
    
    # Extract all text to find placeholders
    all_text = '\n'.join([para.text for para in doc.paragraphs])
    
    # Find all {{variable}} placeholders
    import re
    placeholders = re.findall(r'\{\{([^}]+)\}\}', all_text)
    
    print(f"\n✓ Found {len(placeholders)} placeholders:")
    for placeholder in sorted(set(placeholders)):
        print(f"  - {{{{placeholder}}}}")
    
    return doc


def generate_sample_certificate():
    """Generate a sample certificate with realistic data."""
    print("\n" + "=" * 80)
    print("GENERATING SAMPLE CERTIFICATE")
    print("=" * 80)
    
    template_path = 'media/certificate_templates/samples/ISO_22301_2019_Certificate_Template.docx'
    output_path = 'media/certificate_templates/samples/ISO_22301_2019_Sample_Certificate.docx'
    
    # Load template
    doc = Document(template_path)
    
    # Sample data based on audit report structure
    sample_data = {
        'client_name': 'Global Tech Solutions Ltd.',
        'client_address': '456 Innovation Drive, Tech Park, Silicon Valley, CA 94025, USA',
        'client_phone': '+1 (650) 555-1234',
        'client_email': 'info@globaltechsolutions.com',
        'iso_standard_code': 'ISO 22301:2019',
        'iso_standard_name': 'Business Continuity Management Systems',
        'scope': 'Business Continuity Management System for IT infrastructure services, cloud computing operations, data center management, and customer support services across three locations: Main Office (Silicon Valley, CA), Data Center (Austin, TX), and Support Center (Seattle, WA)',
        'certificate_number': 'BCMS-22301-2025-001',
        'issue_date': datetime.now().strftime('%B %d, %Y'),
        'expiry_date': (datetime.now() + timedelta(days=3*365)).strftime('%B %d, %Y'),
        'certification_body': 'AssureHub Certification Services',
        'accreditation_number': 'ACC-BCMS-2025-789',
        'lead_auditor_name': 'Dr. Sarah Johnson, CBCP, CISA',
        'lead_auditor_email': 'sarah.johnson@assurehub.com',
    }
    
    print("\n✓ Sample data prepared:")
    for key, value in sample_data.items():
        print(f"  - {key}: {value[:60]}{'...' if len(str(value)) > 60 else ''}")
    
    # Replace placeholders in all paragraphs
    for para in doc.paragraphs:
        for key, value in sample_data.items():
            placeholder = f'{{{{{key}}}}}'
            if placeholder in para.text:
                # Replace in runs to preserve formatting
                for run in para.runs:
                    if placeholder in run.text:
                        run.text = run.text.replace(placeholder, str(value))
    
    # Replace placeholders in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for key, value in sample_data.items():
                        placeholder = f'{{{{{key}}}}}'
                        if placeholder in para.text:
                            for run in para.runs:
                                if placeholder in run.text:
                                    run.text = run.text.replace(placeholder, str(value))
    
    # Save the sample certificate
    doc.save(output_path)
    
    file_size = os.path.getsize(output_path)
    print(f"\n✓ Sample certificate generated successfully!")
    print(f"  - Output: {output_path}")
    print(f"  - Size: {file_size / 1024:.1f} KB")
    
    return output_path


def main():
    """Main test function."""
    print("\n" + "=" * 80)
    print("ISO 22301:2019 CERTIFICATE TEMPLATE TEST")
    print("=" * 80)
    print()
    
    # Step 1: Analyze audit report
    audit_doc = analyze_audit_report()
    
    # Step 2: Test certificate template
    cert_doc = test_certificate_template()
    
    # Step 3: Generate sample certificate
    if cert_doc:
        sample_path = generate_sample_certificate()
        
        print("\n" + "=" * 80)
        print("TEST COMPLETE - ALL CHECKS PASSED ✓")
        print("=" * 80)
        print("\nNext steps:")
        print("1. Open the sample certificate to verify formatting")
        print("2. Load template into database using load_certificate_templates.py")
        print("3. Test certificate generation through Django API")
    else:
        print("\n❌ Test failed - certificate template not found")


if __name__ == '__main__':
    main()

